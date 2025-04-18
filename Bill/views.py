from django.shortcuts import render
from django.http import JsonResponse
from Bill.models import *
from Bill.forms import *
from docx import Document
from datetime import datetime
import inflect
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def invoice(request):
    company = Company.objects.all()
    try:
        bill = Bill.objects.all().last().bno    
    except:
        bill = 0
    if request.method == 'POST':
        filledform = InvoiceForm(request.POST)
        total=0
        gt=0
        success=True
        bill_status = ""
        db_status = ""
        uniquebno = ""
        if filledform.is_valid():
            filled_data = {key: value for key, value in filledform.cleaned_data.items() if value}
            success = generatebill(filled_data)
            bill_status = "success" if success else "not success"
            ch_bno = filledform.cleaned_data.get('bno')
            try:
                sample = Bill.objects.get(bno=ch_bno)
                uniquebno = "not unique"
            except:
                uniquebno = "unique"
            try:
                product_to_append = ""
                qty_to_append = ""
                rate_to_append = ""
                amt_to_append = ""
                for i in range(1,5):
                    try:
                        if filled_data[f"product{i}"]:
                            product_to_append += str(filled_data[f"product{i}"])+","
                            qty_to_append += str(filled_data[f"qty{i}"])+","
                            rate_to_append += str(filled_data[f"rate{i}"])+','
                            amt_to_append += str(filled_data[f"amount{i}"])+','
                            total += int(filled_data[f"amount{i}"])
                    except:
                        pass
                gt = total + (total*0.18)
                db_status="not success"
                if filled_data['dbupdate']=="yes" and uniquebno=="unique":
                    bill_id=filledform.save() 
                    Products.objects.create(bill=bill_id,items=product_to_append,qty=qty_to_append,rate=rate_to_append,amt=amt_to_append,hsn='81082000',grandTotal=gt)
                    db_status ="success"
            except Exception as e :
                print(e)
        return render(request, "invoice.html",{"company":company,"no":bill+1,"form":filledform,"subtotal":total,"grand":gt,"gst":F"{total*0.09:.2f}" ,"fileerror":bill_status,"dbstatus":db_status,"ubno":uniquebno})

    return render(request, "invoice.html",{"company":company,"no":bill+1})

def address(request):
    add = Company.objects.get(c_name=request.GET.get('toName')).toAddress
    gst = Company.objects.get(c_name=request.GET.get('toName')).gstIN
    return JsonResponse({'add':add,'gst':gst})

def products(request):
    cName = request.GET.get('toName')
    product = request.GET.get('prod')
    company = Company.objects.get(c_name=cName)
    rate = getattr(company,product)
    return JsonResponse({'price':rate})

def rupess_cov(amount):
    cvt = inflect.engine()
    blk =  int(int(amount)%100000)
    print(str(amount)[0])
    lk = f"{cvt.number_to_words(str(amount)[0])} lakh" if amount//100000 >0 else 0
    rupees = f"{cvt.number_to_words(blk)} Rupees "
    paise = int(str(amount).split(".")[-1]) if len(str(amount).split("."))>=2 else -1
    if paise>0:
        return f"{lk if lk!=0 else ""} {rupees}and {cvt.number_to_words(paise)} Paise only "
    else:
        return  f"{lk if lk!=0 else ""} {rupees}Only"
 
def generatebill(datas):
    items = [
    ]
    try:
        for i in range(1,5):
            if datas[f"product{i}"]:
                product_name = datas[f"product{i}"].split("_")
                product_name = product_name[1:] + [" Mesh"]
                qty = datas[f"qty{i}"]
                rate = datas[f"rate{i}"]
                product = (product_name,"81082000",qty,rate)
                items.append(product)
    except Exception as e:
        print(e)
        pass

    doc = Document("static/Template.docx")
    table = doc.tables[0]  
    date = datas["invoiceDate"].strftime("%d-%m-%Y")
    
    to_bold = table.rows[1].cells[3]
    to_cell = to_bold.paragraphs[0].clear()
    to_cell.add_run("To : ").bold=True
    to_cell.add_run(f"{datas.get('toName', '')},{datas.get('toAdd', '')}")

    table.rows[1].cells[7].text = str(datas.get("bno",""))
    table.rows[2].cells[7].text = date
    gst_bold = table.rows[3].cells[0]
    gst_cell = gst_bold.paragraphs[0].clear()
    gst_cell.add_run("PARTY'S GSTIN NO : ").bold=True
    gst_cell.add_run(datas.get('gstin', ''))
    
    if datas.get('eWayBill'):
        table.rows[3].cells[7].text = datas['eWayBill']
    
    shippingadd_bold = table.rows[4].cells[4]
    shippingAdd_cell = shippingadd_bold.paragraphs[0].clear()
    shippingAdd_cell.add_run(f"Shipping address : ").bold=True
    try:
        if datas["shippingAdd"]:
            shippingAdd_cell.add_run(datas["shippingAdd"])
    except:
        pass

    total_amount = 0
    for idx, (desc, hsn, qty, rate) in enumerate(items, 1):
        row = table.rows[idx + 5]  
        amount = float(qty) * float(rate)
        total_amount += amount
        row.cells[0].text = str(idx)
        print(desc)
        if desc[0] == "scrap":
            row.cells[1].text = "Titanium Scrap"
        else:
            row.cells[1].text = desc+[" Titanium Powder"]
        row.cells[5].text = str(hsn)
        row.cells[6].text = str(qty)+" KG"
        row.cells[8].text = str(rate)+".00"
        row.cells[9].text = f"{amount:.2f}"

        paragraph = row.cells[9].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    table.rows[10].cells[9].text = f"{total_amount:.2f}" # Set horizontal alignment (right)
    paragraph = table.rows[10].cells[9].paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    cgst = total_amount * 0.09
    sgst = total_amount * 0.09
    # igst = total_amount * 0.18
    table.rows[11].cells[9].text = f"{cgst:.2f}"
    table.rows[11].cells[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  

    table.rows[12].cells[9].text = f"{sgst:.2f}"
    table.rows[12].cells[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  

    total_after_tax = total_amount + cgst + sgst
    table.rows[14].cells[9].text = f"{total_after_tax:.2f}"
    table.rows[14].cells[9].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  
    table.rows[15].cells[2].text =  rupess_cov(total_after_tax)
    try:
        doc.save(f"static/invoices/INV_{datas.get('bno', '')}.docx")
        basedir = os.path.dirname(__file__)
        invdir = os.path.dirname(basedir)
        os.startfile(os.path.join(invdir,"static","invoices",f"INV_{datas.get('bno', '')}.docx"))
        return True
    except Exception as e:
        print(e)
    return False